import logging
from typing import Dict, Any

from agent_state import AgentStateDict, estimate_tokens, calculate_context_size
from llm_factory import create_llm
from prompts import get_base_system_prompt
from nodes.schema import SWEPlannerOutput

logger = logging.getLogger(__name__)

# Specialized System Prompt for SWE Agent
SWE_SYSTEM_PROMPT = f"""{get_base_system_prompt()}

ROLE: Expert Software Engineer Agent (SWE).
FOCUS: Code implementation, debugging, refactoring, and file system operations.

AVAILABLE ACTIONS:
- bash: Execute shell commands (git, ls, grep, python scripts, etc.)
- edit: View, create, and modify files.
- complete: Task finished.

GUIDELINES:
1. Prefer 'bash' for exploration (ls, find, grep) and running scripts.
2. Use 'edit' ONLY for viewing or modifying TEXT files (not binary like .docx).
3. For document creation (.docx, .pdf, etc.):
   - Write a Python script inline using bash: echo '...' > script.py && python script.py
   - OR use cat << 'EOF' > script.py ... EOF && python script.py
4. NEVER use 'edit' to create .docx files directly - they are binary!
5. Verify your changes (e.g., ls -la to confirm file creation).

CRITICAL: DETAILED & PEDAGOGICAL CONTENT GENERATION
When creating educational content, documents, courses, or reports:
1. NEVER produce superficial or summary-only content - THIS IS VERY IMPORTANT
2. Each section MUST include:
   - 5+ paragraphs of detailed explanations with context and background
   - Concrete examples, analogies, and real-world applications
   - Technical details (formulas, equations, chemical reactions, diagrams descriptions)
   - Step-by-step breakdowns of complex processes
3. For courses specifically:
   - Learning objectives at the start of each section
   - Key vocabulary/definitions boxes
   - Exercises and review questions at the end
   - Connections between concepts
4. Minimum length: Each major heading should have 500+ words of content
5. Target audience: Students learning this topic for the first time
6. Use clear structure with numbered subheadings, bullet points, and emphasis

QUALITY OVER BREVITY. A comprehensive 20-page document is preferred over a 3-page summary.

EXAMPLE for .docx creation:
For action_details with bash, provide ONLY the shell command(s), not explanations:
cat << 'EOF' > /workspace/create_doc.py
from docx import Document
doc = Document()
doc.add_heading('Title', 0)
doc.add_paragraph('Content')
doc.save('/workspace/output.docx')
EOF
python /workspace/create_doc.py

CRITICAL: In action_details, provide ONLY the executable command or content, never include explanations or numbered steps.
"""

USER_TEMPLATE = """CURRENT CONTEXT:
{context}

STATE:
- Files: {seedbox_manifest}
- Last Output: {last_tool_output}

Iteration: {iteration_count}

What is the next TECHNICAL step? Respond with the SWEPlannerOutput structure."""


def _build_context(state: AgentStateDict) -> str:
    # Simplified context builder for SWE
    messages = state.get("messages", [])
    return "\n".join(
        [
            f"[{m.get('role', '?').upper()}]: {m.get('content', '')}"
            for m in messages[-10:]
        ]
    )  # Focus on recent history


def swe_planner_node(state: AgentStateDict) -> dict:
    """Specialized Planner for SWE tasks using structured output."""
    logger.info(f"SWE Planner - Iteration {state.get('iteration_count', 0)}")

    context = _build_context(state)
    user_message = USER_TEMPLATE.format(
        context=context,
        seedbox_manifest=state.get("seedbox_manifest", [])[:10],
        last_tool_output=str(state.get("last_tool_output", ""))[:2000],
        iteration_count=state.get("iteration_count", 0),
    )

    try:
        llm = create_llm(temperature=0.0)  # Low temp for code precision
        structured_llm = llm.with_structured_output(SWEPlannerOutput)

        messages = [
            {"role": "system", "content": SWE_SYSTEM_PROMPT},
            {"role": "user", "content": user_message},
        ]

        parsed: SWEPlannerOutput = structured_llm.invoke(messages)

        if parsed is None:
            raise ValueError("LLM returned None for structured output")

        logger.debug(
            f"SWE Planner decision: Action={parsed.next_action}, Details={parsed.action_details[:100]}..."
        )

        return {
            "messages": [
                {
                    "role": "assistant",
                    "content": f"Action: {parsed.next_action}\nDetails: {parsed.action_details}",
                }
            ],
            "internal_monologue": parsed.internal_monologue,
            "todo_list": parsed.todo_list,
            "current_action": parsed.next_action,
            "action_details": parsed.action_details,
            "iteration_count": state.get("iteration_count", 0) + 1,
        }
    except Exception as e:
        logger.error(f"SWE Planner crash: {e}")
        return {"current_action": "complete", "action_details": f"Error: {str(e)}"}
