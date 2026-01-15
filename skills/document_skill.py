"""
Document Creation Skill - Generate professional documents for Manus agent.

Provides capabilities to create Word documents (.docx) with structured content
including headings, paragraphs, and lists.
"""

import os
import logging
from typing import Optional, Any
from pathlib import Path

# Configure logging
log_level = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(level=getattr(logging, log_level, logging.INFO))
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Document generation disabled.")


class DocumentSkill:
    """
    Skill for generating structured documents (DOCX).

    Allows creating reports, course materials, and other structured text
    documents with proper formatting.

    Example:
        >>> skill = DocumentSkill()
        >>> await skill.create_word_document(
        ...     "Report.docx",
        ...     content={"title": "My Report", "sections": [...]}
        ... )
    """

    def __init__(self, output_dir: Optional[str] = None) -> None:
        """
        Initialize the document skill.

        Args:
            output_dir: Default output directory.
        """
        self.output_dir = Path(output_dir or os.getcwd())
        logger.debug(f"DocumentSkill initialized (output_dir={self.output_dir})")

    def _ensure_docx_support(self) -> None:
        """Check if python-docx is available."""
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx not installed. Run: uv pip install python-docx"
            )

    async def create_word_document(
        self,
        filename: str,
        content: dict[str, Any],
        overwrite: bool = True,
    ) -> str:
        """
        Create a Word document from structured content.

        Args:
            filename: Name of the output file.
            content: Dict containing document structure:
                - title: str
                - subtitle: Optional[str]
                - sections: List[Dict] usually identifying:
                    - heading: str
                    - content: str or List[str]
            overwrite: Whether to overwrite existing file.

        Returns:
            Absolute path to the created file.

        Example Content Structure:
            {
                "title": "Photosynthesis Course",
                "subtitle": "Introduction to Plant Biology",
                "sections": [
                    {
                        "heading": "Introduction",
                        "content": "Photosynthesis is the process..."
                    },
                    {
                        "heading": "Key Concepts",
                        "content": ["Chlorophyll", "Light Reaction", "Calvin Cycle"]
                    }
                ]
            }
        """
        self._ensure_docx_support()

        # Resolve path
        file_path = self.output_dir / filename
        if not filename.endswith(".docx"):
            file_path = file_path.with_suffix(".docx")

        if file_path.exists() and not overwrite:
            raise FileExistsError(f"File already exists: {file_path}")

        logger.info(f"Creating Word document: {file_path}")

        # Create document
        doc = Document()

        # Add Title
        if title := content.get("title"):
            doc.add_heading(title, 0)

        # Add Subtitle or Description
        if subtitle := content.get("subtitle"):
            p = doc.add_paragraph(subtitle)
            p.italic = True

        # Add Table of Contents (Placeholder text as real TOC is complex in python-docx)
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("auto-generated based on headings below")
        doc.add_page_break()

        # Process Sections
        for section in content.get("sections", []):
            if heading := section.get("heading"):
                doc.add_heading(heading, level=1)

            section_content = section.get("content")

            if isinstance(section_content, list):
                # Bullet points
                for item in section_content:
                    doc.add_paragraph(str(item), style="List Bullet")

            elif isinstance(section_content, str):
                # Standard paragraph
                doc.add_paragraph(section_content)

            # Handle subsections if nested structure exists
            if subsections := section.get("subsections", []):
                for sub in subsections:
                    if sub_head := sub.get("heading"):
                        doc.add_heading(sub_head, level=2)

                    if sub_cont := sub.get("content"):
                        if isinstance(sub_cont, list):
                            for item in sub_cont:
                                doc.add_paragraph(str(item), style="List Bullet")
                        else:
                            doc.add_paragraph(str(sub_cont))

        # Save document
        try:
            doc.save(str(file_path))
            return str(file_path.resolve())
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise RuntimeError(f"Failed to save document: {e}") from e


# ═══════════════════════════════════════════════════════════════════════════════
# STANDALONE TEST
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import asyncio

    async def test_creation():
        skill = DocumentSkill()
        content = {
            "title": "Test Document",
            "subtitle": "Generated by Manus",
            "sections": [
                {
                    "heading": "Introduction",
                    "content": "This is a test paragraph generated automatically.",
                },
                {"heading": "Features", "content": ["Formatting", "Headings", "Lists"]},
            ],
        }
        path = await skill.create_word_document("test_output.docx", content)
        print(f"Created: {path}")

    asyncio.run(test_creation())
